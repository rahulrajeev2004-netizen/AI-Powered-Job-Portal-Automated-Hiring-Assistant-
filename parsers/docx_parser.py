import docx
from utils.logger import get_logger

logger = get_logger('docx_parser', 'logs/extraction.log')

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file handling paragraphs and tables.
    Returns extracted readable text.
    """
    logger.info(f"Extracting text from DOCX: {file_path}")
    text_content = []
    
    try:
        doc = docx.Document(file_path)
        
        # Iterate over all elements (paragraphs and tables) to keep order roughly.
        # But python-docx doesn't easily iterate paragraphs and tables in sequence natively.
        # For simplicity, extract paragraphs first, then tables.
        # A more robust approach iterates through doc.element.body to preserve exact order,
        # but for resume parsing, merging tables often works fine.
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_content.append(text)
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                # Merge row cells into a readable formatted line
                row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if row_data:
                    text_content.append(" | ".join(row_data))
                    
        full_text = "\n".join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from {file_path}")
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""
